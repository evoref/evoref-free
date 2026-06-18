"""DOCX Extractor

.docx ファイルから段落・テーブルのテキストを抽出する。
python-docx を使用。
"""

from __future__ import annotations

from pathlib import Path
from typing import Any, BinaryIO

from backend.extraction._binary_source_base import BinarySourceExtractorBase
from backend.extraction.base import ExtractionError


class DocxExtractor(BinarySourceExtractorBase):
    """DOCX ファイルからテキストを抽出"""

    @property
    def extensions(self) -> frozenset[str]:
        return frozenset({".docx"})

    @property
    def error_code(self) -> str:
        return "docx_error"

    @property
    def requires(self) -> list[str]:
        return ["python-docx"]

    def is_available(self) -> bool:
        try:
            from docx import Document  # noqa: F401
            return True
        except ImportError:
            return False

    def _extract_from_source(
        self,
        source: Path | BinaryIO,
        source_name: str,
    ) -> tuple[str, dict[str, Any]]:
        Document = self._import_docx()
        # python-docx は Path / str / file-like を受け付けるが、
        # 既存挙動 (str(path)) を維持するため Path のときのみ str に変換する
        arg = str(source) if isinstance(source, Path) else source
        doc = Document(arg)
        text, para_count, table_count = self._extract_content(doc)
        return text, {"paragraphs": para_count, "tables": table_count}

    @staticmethod
    def _import_docx():
        try:
            from docx import Document
            return Document
        except ImportError:
            raise ExtractionError(
                "missing_library",
                "python-docx is required for DOCX extraction: pip install python-docx",
            )

    @staticmethod
    def _extract_content(doc) -> tuple[str, int, int]:
        """段落とテーブルからテキストを抽出"""
        lines: list[str] = []
        para_count = 0
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text)
                para_count += 1
        for table in doc.tables:
            for row in table.rows:
                cells = "\t".join(cell.text.strip() for cell in row.cells)
                if cells.strip():
                    lines.append(cells)
        return "\n".join(lines), para_count, len(doc.tables)
