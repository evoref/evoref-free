"""DOCX Writer

.docx: ContentBlock → Word 文書（見出し・段落・表・コードブロック）
python-docx を使用。
"""

from __future__ import annotations

import io
import re

from backend.export._writer_base import BytesWriterBase
from backend.export.base import ExportContent


def _add_inline_runs(paragraph, text: str) -> None:
    """inline Markdown を解析して Word の Run に変換"""
    # 簡易パース: bold, italic, code を検出
    pattern = re.compile(
        r"(\*\*(.+?)\*\*|__(.+?)__)"    # bold
        r"|(\*(.+?)\*|_(.+?)_)"          # italic
        r"|(`(.+?)`)"                     # code
    )
    last_end = 0
    for m in pattern.finditer(text):
        # マッチ前の通常テキスト
        if m.start() > last_end:
            paragraph.add_run(text[last_end:m.start()])

        if m.group(2) or m.group(3):
            # bold
            run = paragraph.add_run(m.group(2) or m.group(3))
            run.bold = True
        elif m.group(5) or m.group(6):
            # italic
            run = paragraph.add_run(m.group(5) or m.group(6))
            run.italic = True
        elif m.group(8):
            # code
            run = paragraph.add_run(m.group(8))
            run.font.name = "Consolas"
        last_end = m.end()

    # 残りのテキスト
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


def _build_docx(content: ExportContent) -> bytes:
    """ExportContent を DOCX バイトデータに変換"""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()

    # デフォルトフォント設定
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)

    blocks = content.blocks
    if not blocks and content.raw_markdown:
        from backend.export.content_converter import ContentConverter
        blocks = ContentConverter().convert(content.raw_markdown)

    for block in blocks:
        if block.type == "heading":
            doc.add_heading(block.content, level=block.level)

        elif block.type == "paragraph":
            para = doc.add_paragraph()
            _add_inline_runs(para, block.content)

        elif block.type == "code":
            para = doc.add_paragraph()
            run = para.add_run(block.content)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            # グレー背景のシェーディングは python-docx で直接サポートされないため省略

        elif block.type == "table":
            if block.rows:
                table = doc.add_table(
                    rows=len(block.rows),
                    cols=len(block.rows[0]),
                    style="Table Grid",
                )
                for i, row_data in enumerate(block.rows):
                    for j, cell_text in enumerate(row_data):
                        table.rows[i].cells[j].text = cell_text
                    # ヘッダー行を太字に
                    if i == 0:
                        for cell in table.rows[0].cells:
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    run.bold = True

        elif block.type == "list":
            for idx, item in enumerate(block.items):
                style_name = "List Number" if block.ordered else "List Bullet"
                para = doc.add_paragraph(style=style_name)
                _add_inline_runs(para, item)

        elif block.type == "quote":
            para = doc.add_paragraph(block.content)
            para.style = doc.styles["Quote"] if "Quote" in [s.name for s in doc.styles] else None
            if para.style is None:
                para.paragraph_format.left_indent = Pt(36)

        elif block.type == "hr":
            doc.add_paragraph("_" * 50)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class DocxWriter(BytesWriterBase):
    """DOCX ファイル Writer"""

    @property
    def extensions(self) -> frozenset[str]:
        return frozenset({".docx"})

    @property
    def requires(self) -> list[str]:
        return ["python-docx"]

    def _render_bytes(self, content: ExportContent, ext: str) -> bytes:  # noqa: ARG002
        return _build_docx(content)

    def is_available(self) -> bool:
        try:
            from docx import Document  # noqa: F401
            return True
        except ImportError:
            return False
